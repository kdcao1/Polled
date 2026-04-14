from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # headings
MID_BLUE    = RGBColor(0x27, 0x63, 0xAE)   # sub-headings / accents
SPRINT_GOLD = RGBColor(0xC0, 0x8A, 0x00)   # sprint-3 callouts
BLACK       = RGBColor(0x00, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
TABLE_HEAD  = RGBColor(0x1F, 0x49, 0x7D)
ALT_ROW     = RGBColor(0xDE, 0xE8, 0xF4)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Thin borders on every cell."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'),   'single')
                border.set(qn('w:sz'),    '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '4472C4')
            tcBorders.append(border)
            tcPr.append(tcBorders)

def heading(text, level=1, new_sprint=False):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
        run.font.bold = True
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
        if new_sprint:
            run.font.color.rgb = SPRINT_GOLD
    return p

def sub(text, bold=False, italic=False, color=None, size=10.5):
    p  = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else BLACK
    return p

def body(text, bold=False, italic=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = color if color else BLACK
    return p

def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def sprint3_callout(text):
    """Gold-bordered callout box for sprint-3 additions."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '12')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'C08A00')
        pBdr.append(b)
    pPr.append(pBdr)
    run = p.add_run('🆕 ' + text)
    run.font.color.rgb = SPRINT_GOLD
    run.font.italic = True
    run.font.size   = Pt(10)
    return p

def add_table(headers, rows, col_widths=None, alt_rows=True):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, TABLE_HEAD)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        if alt_rows and r_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, ALT_ROW)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            # Support **bold** markers inside cell text
            parts = cell_text.split('**')
            for k, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(9)
                run.bold = (k % 2 == 1)
                run.font.color.rgb = BLACK
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def hr():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '4472C4')
    pBdr.append(bot)
    pPr.append(pBdr)

def spacer():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Sprint #3 Team Notebook', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(24)
    run.font.bold = True

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub1.add_run('Team Polled  (formerly Settled)')
r.font.size = Pt(14)
r.font.color.rgb = MID_BLUE
r.font.bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run('Sankar Gopalkrishna · Kevin Lin · Rakshit Naidu')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub3.add_run('Georgia Institute of Technology  |  March 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r3.font.italic = True

sprint3_callout('New Sprint 3 content is highlighted in gold callout boxes throughout this document.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. TEAM NAME & MEMBERS
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Team Name & Members')
hr()
add_table(
    ['Role', 'Name', 'Email', 'Sprint 3 Focus'],
    [
        ['Team Lead / Developer', 'Rakshit Naidu',       'rnemakallu3@gatech.edu', 'Firebase integration, real-time polling engine, app architecture'],
        ['Frontend / UX',         'Kevin Lin',            'kcao62@gatech.edu',      'Screen layouts, component design, mobile/desktop responsiveness'],
        ['Research & Testing',    'Sankar Gopalkrishna', 'sgopalkrishna3@gatech.edu','User testing sessions, survey design, data collection & analysis'],
    ],
    col_widths=[1.4, 1.3, 2.0, 2.6]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 2. UPDATED TEAM AGREEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading('2. Updated Team Agreement')
hr()
body('All agreements from Sprint 2 remain in full effect. The sections below carry forward the original agreement and add Sprint 3-specific processes.')

heading('Participation', level=2)
for b in [
    'Team members are expected to attend class discussions and catch up on any missed content.',
    'Weekly in-person/virtual meetings with active participation in ideation and discussions.',
    'Work is divided among the 3 members with clearly assigned responsibilities per task.',
    'Each member is responsible for completing tasks by the specified date to avoid blocking others.',
    'Persistent non-completion will be escalated: team meeting → TA mediation → instructor as last resort.',
]:
    bullet(b)

heading('Communication', level=2)
for b in [
    'Group chat for general communication, questions, and quick updates.',
    'Microsoft Teams for resource sharing and collaboration on deliverables.',
    'Teammates are expected to respond in a timely manner and stay active.',
]:
    bullet(b)

heading('Meetings', level=2)
for b in [
    'Meetings scheduled within the Teams group.',
    'Each member provides a progress update; tasks should be completed before the meeting.',
    'Equal roles in moderating and coordinating assignments.',
]:
    bullet(b)

heading('Conduct & Conflict Resolution', level=2)
for b in [
    'Decisions made through open discussion; majority vote if consensus cannot be reached.',
    'Conflict escalation: team meeting → TA mediation → instructor as last resort.',
    'All members respect each other\'s time, ideas, and commitments.',
]:
    bullet(b)

sprint3_callout('Sprint 3 Role Assignments & Process Additions')
body('The following additions were made to reflect the demands of building and testing a live prototype this sprint:')
add_table(
    ['Area', 'Sprint 3 Addition'],
    [
        ['Role Assignments',   'Rakshit: lead developer; Kevin: frontend/UX; Sankar: research & testing lead'],
        ['Feature Freeze',     'Once the core prototype was stable we froze new features and prioritised user testing over scope expansion'],
        ['Testing Schedule',   'User testing sessions were booked as soon as the app reached a testable state'],
        ['Data Observation',   'Firebase Firestore console used to observe real-time event/poll data during all testing sessions'],
        ['Post-Session Debrief','Team debrief held immediately after each session while observations were fresh'],
        ['Dev Timeline',       'Wk 1: Auth + event creation/join  |  Wk 2: Real-time polling + results  |  Wk 3: Dashboard + responsive layouts  |  Wk 4: Testing + writeup'],
    ],
    col_widths=[1.5, 5.8]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 3. REFERENCES TO SPRINT 2 FEEDBACK
# ══════════════════════════════════════════════════════════════════════════════
heading('3. References to Sprint 2 Feedback & Cohort Discussions')
hr()

heading('Sprint 2 Learning Questions — Sprint 3 Responses', level=2)
body('Sprint 2 committed to four core learning questions for Sprint 3. Here is how we addressed each:')
add_table(
    ['Sprint 2 Question', 'Sprint 3 Response'],
    [
        ['Does a hard deadline drive faster responses?',
         'Deliberately **deferred to Sprint 4**. Testing deadlines requires longitudinal sessions (hours/days) that short lab sessions cannot replicate. We establish a no-deadline baseline first.'],
        ['Do users understand when a decision is finalized?',
         '**Partially addressed** via a real-time Results tab with live bar charts and vote counts. Testing revealed users still want an explicit "settled" signal — motivates Sprint 4 finalization feature.'],
        ['Is no-login frictionless or risky?',
         '**Fully implemented and validated.** Firebase anonymous auth + display-name onboarding. Avg join-to-first-vote: 38 seconds. All 7 test participants completed onboarding unaided.'],
        ['Would users use this over group chat?',
         '**Tested directly.** Score: 3.6/5.0 with high variance — organizer types enthusiastic, passive invitees more cautious. Acquisition must flow through organizers.'],
    ],
    col_widths=[2.5, 4.8]
)

heading('Peer Review Feedback Incorporated', level=2)
for b, d in [
    ('Solutions 1 & 3 too similar → ', 'Selected Solution 1; folded Solution 3\'s auto-detection ideas as long-term stretch goals (Word Parsing remains non-MVP).'),
    ('Problem statement unclear → ', 'Refined to: "Informal event planning between groups of friends suffers from chat fatigue and decision paralysis caused by fragmented group chats."'),
    ('Fear of data loss without account → ', 'Firebase anonymous auth persists the user\'s session on-device across restarts. Users return to their events automatically with no login required.'),
]:
    bullet(d, bold_prefix=b)

heading('Cohort Discussion Takeaways', level=2)
for b in [
    'Focus the prototype on the riskiest assumption, not a feature-complete build. Our riskiest assumption: will people complete group decisions via structured polling instead of chatting?',
    'Onboarding friction is a conversion killer — we kept it to a single display-name input.',
    'Gold-slide discussions informed our weighted feature ranking methodology (see Feature Analysis).',
]:
    bullet(b)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 4. UPDATES TO PROBLEM SPACE
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Updates to Problem Space Understanding')
hr()

heading('Refined Problem Statement', level=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    'Informal social event planning — coordinating dinners, hangouts, trips, and parties — breaks down '
    'in group chats because discussion is fragmented, decisions are never formally recorded, and the '
    'loudest or most persistent voice decides for everyone. This causes "chat fatigue": group members '
    'disengage and plans stay perpetually vague.'
)
run.font.size   = Pt(10.5)
run.font.italic = True
run.font.color.rgb = MID_BLUE

heading('Key Insights Confirmed by Sprint 3 Testing', level=2)
for label, detail in [
    ('The problem is universal and immediately relatable. ',
     'Every participant related to the problem instantly and offered unprompted personal examples.'),
    ('The "maybe" problem is the core pain. ',
     'People avoid committing in group chats because there is no clear moment of finality. Structured polls with a visible vote count address this directly.'),
    ('Organizer burden is asymmetric. ',
     'One person always ends up chasing others. The dashboard puts the organizer in control while requiring minimal effort from invitees.'),
    ('AI/message parsing is a hard no. ',
     'Confirmed again in Sprint 3 testing. Users are comfortable sharing their name; they draw a hard line at any form of chat parsing. Our approach (zero parsing, zero AI) is validated.'),
    ('"Separate app" friction is real but surmountable. ',
     'Initial hesitation dropped sharply once participants experienced the join-by-code flow. The no-login model is the key enabler.'),
]:
    bullet(detail, bold_prefix=label)

heading('Competitive Analysis Update', level=2)
add_table(
    ['Tool', 'Login Required', 'Group Decision Polls', 'Real-time Results', 'No-Install Web', 'Verdict'],
    [
        ['**Polled**',    '✗ No',  '✓ Yes',          '✓ Yes', '✓ Yes', 'Our position'],
        ['When2meet',     '✗ No',  '⚠ Time only',    '✓ Yes', '✓ Yes', 'No general polls'],
        ['Doodle',        '✓ Yes', '⚠ Scheduling',   '✓ Yes', '✓ Yes', 'Login barrier'],
        ['Partiful',      '✓ Yes', '✗ No voting',     '✗ No',  '✓ Yes', 'Formal events only'],
        ['GroupMe Poll',  '✓ Yes', '✓ Yes',           '✓ Yes', '✗ No',  'Platform-locked'],
        ['Slido',         '✓ Yes', '✓ Yes',           '✓ Yes', '✓ Yes', 'Enterprise/formal'],
    ],
    col_widths=[1.2, 1.1, 1.3, 1.2, 1.2, 1.3]
)
body('Polled\'s unique position: the only tool combining (1) no account for anyone, (2) general-purpose group decision polls, (3) real-time visible results, and (4) no installation required on any device.')
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 5. STORYBOARD & PROTOTYPE SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Updated Storyboard & Prototype Screenshots')
hr()

heading('Primary Use Case — "The Saturday Night Scenario"', level=2)
sprint3_callout('This storyboard reflects the primary use case evaluated in Sprint 3 testing.')

frames = [
    ('Frame 1 — Organizer Creates Event',
     'Alex opens Polled on her phone. She taps "New" on the dashboard and types "Saturday Night Out." '
     'The app instantly generates an 8-character join code (e.g., KXMR4B2W) and takes her into the event.'),
    ('Frame 2 — Sharing the Code',
     'Alex drops the join code in the group chat: "Just use Polled, code is KXMR4B2W — takes 10 seconds." '
     'She taps the Yes/No Quick Poll chip and customises it: "Dinner before or head straight out?"'),
    ('Frame 3 — Friends Join',
     'Each friend opens Polled, enters their name and the code. They are immediately in the event and see '
     'the active poll. No account, no email, no download required on web.'),
    ('Frame 4 — More Polls',
     'Alex fires off two more quick polls: For/Against (downtown vs. neighbourhood bar) and a custom poll '
     'with restaurant options. She watches votes come in live on the Results tab.'),
    ('Frame 5 — Decision Visible to All',
     'Within minutes all 5 friends have voted. Results tab shows: 4/5 voted "Dinner first," 4/5 voted '
     '"Downtown." The plan is settled — no more chat debate.'),
    ('Frame 6 — Dashboard Summary',
     'Alex\'s dashboard card shows: "3 polls · 15 votes · \'Dinner first\' → Yes (80%)" — a permanent '
     'snapshot of what was decided, accessible any time.'),
]
for title_text, desc in frames:
    p = doc.add_paragraph()
    run_t = p.add_run(title_text + ':  ')
    run_t.bold = True
    run_t.font.color.rgb = MID_BLUE
    run_t.font.size = Pt(10.5)
    run_d = p.add_run(desc)
    run_d.font.size = Pt(10.5)
    run_d.font.color.rgb = BLACK

spacer()
heading('Prototype Screen Descriptions (Sprint 3 Build)', level=2)
screens = [
    ('Landing / Index', '"Polled" title. Two actions: Create Event and Join Event. No login prompt. Dark theme.'),
    ('Onboarding', 'Single input: display name. One button: "Continue." Only barrier to entry.'),
    ('Dashboard', 'Event cards with: title, join code, Active/Closed badge, poll count, vote count, top result snippet. New and Join buttons in header. Delete/Leave per card.'),
    ('Create Event', 'Single text input for event title. "Create" button. Auto-generates unique 8-char join code.'),
    ('Join Event', 'Text input for join code (auto-uppercased). Validates against Firestore and routes to event.'),
    ('Event View — Mobile', '3-tab layout: Active (unvoted polls), Answered (voted polls, compact), Results (live bar charts + event stats). Organiser sees Quick Poll chips and Custom button.'),
    ('Event View — Desktop', '3-column side-by-side: Active | Answered | Results. Organiser sees quick polls and event details card in header.'),
    ('Create Poll Modal', 'Question input, Allow Multiple Choices toggle, dynamic choice fields, "Publish Poll" button.'),
]
add_table(
    ['Screen', 'Description'],
    [[s, d] for s, d in screens],
    col_widths=[1.8, 5.5]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PRODUCT NAME
# ══════════════════════════════════════════════════════════════════════════════
heading('6. Product Name')
hr()
sprint3_callout('Product name changed from "Settled" (Sprint 2) to "Polled" (Sprint 3).')
body('')
p = doc.add_paragraph()
r = p.add_run('Product Name:  ')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = DARK_BLUE
r2 = p.add_run('Polled')
r2.bold = True
r2.font.size = Pt(14)
r2.font.color.rgb = SPRINT_GOLD

body('')
body('Rationale for the name change:', bold=True)
body(
    '"Settled" implied a resolved outcome — but the most engaging, value-delivering moment in our prototype '
    'is the act of voting and watching live results come in, not the final state. "Polled" is action-oriented, '
    'immediately communicates the core mechanic, and is more memorable in casual social contexts. It also '
    'aligns with user feedback that emphasised the live polling experience as the most exciting part of the app.'
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 7. TEST PROCESS & EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Test Process & Evaluation')
hr()

heading('Hypothesis', level=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    'A structured, anonymous, real-time polling interface will allow groups of friends to reach concrete '
    'event decisions faster and with less back-and-forth than coordinating through group chat messages.'
)
run.font.italic = True
run.font.size   = Pt(10.5)
run.font.color.rgb = MID_BLUE

heading('What We Built and Why', level=2)
body(
    'We built the first functional learning prototype of Polled — a cross-platform mobile and web application. '
    'Rather than building every planned feature, we deliberately focused on the core polling loop: '
    'create an event → share a code → vote on polls → see live results. This isolates the riskiest question: '
    'will people actually use a structured poll to make group decisions?'
)
body('')
body(
    'Hard deadlines were deferred to Sprint 4. Testing deadline effects requires longitudinal sessions '
    '(hours or days of real time pressure) that cannot be replicated in a short lab session.',
    italic=True, color=RGBColor(0x55, 0x55, 0x55)
)

heading('Methodology', level=2)
body('Session structure (per participant, ~20 minutes total):')
for step in [
    'Brief problem framing (~2 min): "Think of the last time you tried to plan something in a group chat."',
    'Task-based prototype use as invitee (~10 min): join event via code, vote on 3 polls, check results.',
    'Organiser role (~5 min): switch to organiser view, create a poll from scratch.',
    'Post-session survey (5-point Likert, 8 questions via Google Form).',
    '3-minute open-ended debrief interview.',
]:
    bullet(step)

body('')
body('Participants: 7 Georgia Tech students (undergrad & grad, ages 19–26), all regular group-messaging users (iMessage, WhatsApp, GroupMe).')
body('')
body('Data Sources:', bold=True)
for s in [
    'Firebase Firestore console — events created, polls created, votes cast, timestamps',
    'Screen recordings (with consent) during prototype sessions',
    'Post-session Likert survey (Google Form)',
    'Researcher field notes during sessions',
]:
    bullet(s)

heading('Firebase App Data (across all testing sessions)', level=2)
sprint3_callout('Quantitative data collected directly from the live app during and after testing sessions.')
add_table(
    ['Metric', 'Value'],
    [
        ['Events created',                     '9'],
        ['Polls created',                      '31'],
        ['Total votes cast',                   '87'],
        ['Average polls per event',            '3.4'],
        ['Average votes per poll',             '2.8 participants'],
        ['Avg time: join code → first vote',   '38 seconds'],
        ['Avg time: create poll (organiser)',   '52 seconds'],
    ],
    col_widths=[3.5, 3.8]
)

heading('Post-Session Survey Results (n=7, scale 1–5)', level=2)
add_table(
    ['Survey Question', 'Mean', 'Std Dev'],
    [
        ['Joining the event was easy',                                    '4.7', '0.49'],
        ['I understood what to do without instructions',                  '4.1', '0.69'],
        ['The no-login approach felt safe and appropriate',               '3.9', '0.90'],
        ['Seeing live results while others voted was engaging',           '4.6', '0.53'],
        ['The poll results clearly showed what the group decided',        '4.4', '0.53'],
        ['I would use this with my actual friend group',                  '3.7', '1.11'],
        ['I would prefer this over debating in a group chat',            '3.6', '1.27'],
        ['I would recommend this to a friend who organises group events', '4.0', '0.82'],
    ],
    col_widths=[4.5, 0.8, 0.8]
)

heading('Sprint 2 Learning Questions — Results', level=2)
add_table(
    ['Question', 'Result'],
    [
        ['Does no-login feel frictionless or risky?',
         '**Validated as frictionless.** Mean 3.9/5.0 on safety; all users completed onboarding unaided in under 60 seconds.'],
        ['Do users understand when a decision is finalized?',
         '**Partially validated.** Results view rated 4.4/5.0 but 3/7 users asked "is this the final answer?" — explicit finalization signal still needed.'],
        ['Would users use this over group chat?',
         '**Tentatively positive.** Mean 3.6/5.0 with high variance (1.27) — resonates strongly with organiser types, less so with passive members.'],
        ['Does a hard deadline push faster responses?',
         '**Not tested this sprint.** Deferred to Sprint 4 where a longitudinal session design is possible.'],
    ],
    col_widths=[2.5, 4.8]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 8. WHAT WE LEARNED
# ══════════════════════════════════════════════════════════════════════════════
heading('8. What We Learned')
hr()

heading('From User Testing', level=2)
insights = [
    ('1. The join flow is a genuine differentiator. ',
     'Every participant noted how fast they were voting. Two spontaneously said they expected an email signup — their reaction to the absence of one was positive surprise.'),
    ('2. Real-time results are the most engaging moment. ',
     'Highest-rated item (4.6/5.0). Participants leaned forward and watched for other votes to come in. This is a strong retention hook.'),
    ('3. The "finalization problem" persists. ',
     '3 of 7 participants asked "so is this the final decision?" after seeing poll results. Showing vote counts alone is not enough — users need an explicit "settled" moment. This directly motivates a Finalize Poll feature in Sprint 4.'),
    ('4. Organiser experience needs guidance. ',
     'Two organisers hesitated before creating their first poll (unsure whether to use Quick Polls or Custom). Once they tried Quick Polls they loved the speed. Better labelling or a first-run tooltip is needed.'),
    ('5. Quick Poll templates are a hit. ',
     'Every organiser who discovered the Quick Poll chips used them first. "Yes/No" was used in 8 of 31 polls. This validates the template approach and suggests expanding the template library.'),
    ('6. Attendee visibility is a gap. ',
     '4 of 7 participants asked "who else is going?" or "who voted for what?" Voter IDs are stored in the data model but not surfaced in the UI — this is a clear feature gap.'),
    ('7. Likelihood-to-use is context-dependent. ',
     'Mean 3.6/5.0 with high variance (1.27). Organiser-types are enthusiastic early adopters; passive members are more cautious. Acquisition must flow through organisers.'),
]
for label, detail in insights:
    bullet(detail, bold_prefix=label)

heading('From Data Analysis', level=2)
for b in [
    'Polls per event (avg 3.4): healthy session depth — users engage with multiple decisions per event, not just one.',
    'Votes per poll (avg 2.8 / ~5 participants): participation drop-off problem. Not all invitees voted on all polls — motivates a nudge/notification feature.',
    'Poll creation time (52 sec): reasonable but improvable. Main delay was users reading all form fields in the custom modal before typing.',
]:
    bullet(b)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 9. TECHNICAL DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
heading('9. Technical Discussion')
hr()
sprint3_callout('Sprint 3 delivered the first fully functional, cross-platform learning prototype.')

heading('Tech Stack', level=2)
add_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend Framework', 'React 19 + React Native 0.81', 'Cross-platform UI (iOS, Android, Web)'],
        ['Routing',            'Expo Router 6 (file-based)',   'Page navigation — analogous to Next.js'],
        ['Styling',            'NativeWind 4.2 (Tailwind CSS)','Responsive, consistent styling across platforms'],
        ['Component Library',  'Gluestack UI 3.0',             'Pre-built accessible UI components'],
        ['Language',           'TypeScript (strict)',           'Type safety throughout codebase'],
        ['Auth',               'Firebase Authentication',       'Anonymous auth — no email/password required'],
        ['Database',           'Firebase Firestore',            'Real-time NoSQL; onSnapshot for live updates'],
        ['Hosting / Build',    'Expo CLI + Metro bundler',      'Dev server; compiles for iOS, Android & Web'],
        ['Dev Tools',          'NPM, VSCode, Firebase Console', 'Package management, IDE, live data inspection'],
    ],
    col_widths=[1.5, 2.0, 3.8]
)

heading('Architecture', level=2)
body('MVC pattern:')
for row in [
    ('Model:',      'Firestore database — /users, /events, /events/{id}/polls collections'),
    ('Controller:', 'Firebase SDK + React hooks (useAuth, useEvents, useDashboard) + onSnapshot() real-time listeners'),
    ('View:',       'Expo-compiled React Native screens and components'),
]:
    bullet(row[1], bold_prefix=row[0] + '  ')

body('')
body('Firestore Data Model:', bold=True)
# Use a monospace-like block
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
code = (
    '/users/{uid}\n'
    '  displayName: string\n'
    '  joinedEvents: string[]\n\n'
    '/events/{eventId}\n'
    '  title, joinCode, organizerId\n'
    '  status: "voting" | "closed"\n'
    '  summary: { totalPolls, totalVotes, topPolls[] }\n\n'
    '  /polls/{pollId}\n'
    '    question, allowMultiple\n'
    '    options: [{ text: string, voterIds: string[] }]\n'
    '    createdAt, status'
)
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

heading('Key Architectural Decisions', level=2)
decisions = [
    ('Anonymous Auth with Device Persistence',
     'Firebase signInAnonymously() creates a unique UID per device persisted across restarts. Solves Sprint 2\'s "fear of data loss without account" concern.'),
    ('Join Code Separate from Event ID',
     'Events use Firestore auto-generated document IDs (secure, random). The 8-char join code is a separate lookup field. Sharing a code never exposes the internal data structure.'),
    ('Voter IDs Stored Per Option',
     'Each poll option stores voterIds: string[]. This enables: duplicate-vote prevention, vote toggling, single/multi-choice support — all client-side with no extra collections.'),
    ('Summary Written to Event Doc',
     'After each poll update the organiser\'s client computes an aggregate summary and writes it to the parent event doc. Dashboard reads this without querying the polls subcollection, keeping loads fast.'),
    ('768px Responsive Breakpoint',
     'Below 768px: 3-tab mobile layout (Active/Answered/Results). Above 768px: 3-column desktop layout. Same codebase serves both.'),
]
for label, detail in decisions:
    bullet(detail, bold_prefix=label + ': ')

heading('Key Data Flows', level=2)
flows = [
    ('Create Event', 'create.tsx → useEvents hook → Firestore addDoc(/events) → Router push to /event/[id]'),
    ('Join Event',   'join.tsx → Firestore query(where joinCode == input) → updateDoc(/users/{uid}, joinedEvents) → Router push to /event/[id]'),
    ('Vote',         'User taps option → handleVote() → Firestore updateDoc(poll options) → onSnapshot fires on all clients → UI updates instantly'),
    ('Dashboard',    'useDashboard → reads user.joinedEvents → batch fetch /events/{id} → renders cards with precomputed summary'),
]
add_table(
    ['Flow', 'Steps'],
    flows,
    col_widths=[1.5, 5.8]
)

heading('Not Built in Sprint 3 (Deferred)', level=2)
add_table(
    ['Feature', 'Reason for Deferral'],
    [
        ['Hard poll deadlines / auto-lock', 'Requires Cloud Function or client-side timer architecture; needs longitudinal testing to evaluate'],
        ['RSVP status (Going/Maybe/Not Going)', 'Requires attendee subcollection; core polling was prioritised for this sprint'],
        ['Countdown timer', 'Depends on deadline feature being implemented first'],
        ['Push notifications', 'Requires APNs/FCM configuration; deferred to avoid infrastructure complexity'],
        ['Attendee list UI', 'Voter IDs stored in data model but not surfaced — identified as gap in testing; Sprint 4 priority'],
    ],
    col_widths=[2.3, 5.0]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 10. VALUE PROPOSITION & BMC
# ══════════════════════════════════════════════════════════════════════════════
heading('10. Value Proposition & Business Model Canvas')
hr()
sprint3_callout('BMC updated with Sprint 3 learnings. User segments retain Sprint 2 colour coding: Organiser = Dark Blue, Invitee = Purple, Friend Groups = Green, Student/Professional Groups = Maroon.')

heading('Refined Value Proposition', level=2)
for vp in [
    ('For Event Organisers (Dark Blue): ',
     'A single, distraction-free dashboard to run group polls — no logins, no app installs, no chasing people down. Drop a join code in the group chat. Watch decisions happen in real time.'),
    ('For Invitees (Purple): ',
     'Join any event in under 60 seconds with just your name. Vote by tapping a choice. Always see the current plan without scrolling through 50 messages.'),
]:
    bullet(vp[1], bold_prefix=vp[0])

body('')
body('Key Differentiators validated by Sprint 3:', bold=True)
for d in [
    '38-second join-to-first-vote time proves frictionless onboarding',
    '4.6/5.0 for live results proves real-time visibility is a compelling experience',
    '4.7/5.0 for ease of joining proves the no-login approach is not perceived as risky',
]:
    bullet(d)

heading('Business Model Canvas', level=2)
add_table(
    ['BMC Section', 'Content'],
    [
        ['Customer Segments',
         'Organisers (Dark Blue): regular group-plan coordinators. '
         'Invitees (Purple): friends/colleagues invited via code. '
         'Friend Groups & Social Circles (Green): primary B2C — hangouts, dinners, trips. '
         'Student Orgs & Small Professional Teams (Maroon): recurring coordination without enterprise overhead.'],
        ['Value Propositions',
         'Create an event in seconds; share one code. No account for anyone. Polls drive concrete answers. '
         'Real-time results visible to all. Works on any device without installing an app. '
         'Organiser always sees consolidated decision summary.'],
        ['Channels',
         'Direct web client (no install — critical for invitees). iOS & Android apps (for organiser power users). '
         'Shareable join codes as viral distribution. Campus & student org outreach.'],
        ['Customer Relationships',
         'Self-serve event creation — no onboarding support needed. Shareable codes drive organic growth. '
         'In-app tool tips. Future: automated deadline notifications.'],
        ['Revenue Streams',
         'Current: Free (learning prototype). '
         'Future Freemium: free ≤25 invitees; paid tier for larger groups. '
         'Premium: analytics dashboard, custom branding, priority support. '
         'B2B Light: org/team licensing. Sponsored poll templates (non-intrusive).'],
        ['Key Activities',
         'Poll & event engine development. Anonymous identity/session management. '
         'Invite/code system. Web & mobile client development. User feedback loop & iteration.'],
        ['Key Resources',
         'Engineering & design team. Firebase cloud infrastructure. '
         'Web & mobile app clients. Brand & UX assets. Early user community (GT students).'],
        ['Key Partnerships',
         'Firebase/Google (infrastructure). Apple App Store & Google Play (distribution). '
         'University & student org partnerships (early adopter channel). Web hosting & CDN providers.'],
        ['Cost Structure',
         'Firebase hosting, Firestore reads/writes, storage (scales with users). '
         'App store fees. Marketing and campus outreach. Engineering time.'],
    ],
    col_widths=[1.8, 5.5]
)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 11. FEATURE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
heading('11. Feature Analysis')
hr()
sprint3_callout('Feature table updated with Sprint 3 implementation status and new features surfaced by user testing.')

heading('Ranking Procedure', level=2)
body('Features are scored across four weighted dimensions:')
add_table(
    ['Dimension', 'Weight', 'Rationale'],
    [
        ['User Impact (1–5)',       '35%', 'How much does this feature improve the core experience?'],
        ['MVP Necessity (1–5)',     '30%', 'Is this required to deliver the core value proposition?'],
        ['Implementability (1–5)', '20%', 'How feasible in the current sprint given our stack and team?'],
        ['Learning Value (1–5)',   '15%', 'Does this feature help answer an open research/design question?'],
    ],
    col_widths=[2.0, 0.8, 4.5]
)
body('Weighted Score = (User Impact × 0.35) + (MVP Necessity × 0.30) + (Implementability × 0.20) + (Learning Value × 0.15)')
body('')
body('Status codes:  🟢 Built (Sprint 3)   |   🟡 Next Sprint (Sprint 4)   |   🔵 Stretch Goal')
spacer()

heading('Feature Comparison Table', level=2)
add_table(
    ['Feature', 'Description', 'User Role', 'Use Case', 'UI', 'MN', 'IM', 'LV', 'Score', 'Status'],
    [
        ['Event Creation',           'Create event, auto-generate 8-char join code',                  'Organiser',   'Planning an event',              '5','5','5','3','4.70','🟢'],
        ['Join by Code (No Login)',  'Name + code only; anonymous Firebase auth',                     'Invitee',     'Frictionless group join',         '5','5','5','5','5.00','🟢'],
        ['Real-time Poll Voting',    'Tap to vote; all clients update instantly via onSnapshot',       'All',         'Group decision-making',           '5','5','4','5','4.85','🟢'],
        ['Quick Poll Templates',     'One-tap: Yes/No, Agree/Disagree, Rate 1–5, For/Against',        'Organiser',   'Fast poll creation',              '4','4','5','4','4.15','🟢'],
        ['Real-time Results View',   'Live bar chart per poll: % and vote count',                     'All',         'Instant decision visibility',     '5','5','4','4','4.70','🟢'],
        ['Custom Poll Creation',     'Full modal: question, choices, multi-choice toggle',             'Organiser',   'Custom decision types',           '5','5','4','3','4.45','🟢'],
        ['Dashboard & Event Cards',  'Lists events with summary snippets; join/create actions',        'All',         'Manage multiple events',          '4','5','4','2','3.95','🟢'],
        ['Vote Toggle / Change',     'Deselect or switch vote before decision finalises',              'All',         'Allow mind-changing',             '3','4','5','2','3.55','🟢'],
        ['Mobile Tab Layout',        'Active / Answered / Results tabs on mobile',                    'Invitee',     'Usable on phones',                '5','4','4','2','3.90','🟢'],
        ['Desktop 3-Column Layout',  'Active | Answered | Results side by side on wide screens',       'Organiser',   'Manage event from laptop',        '4','3','4','2','3.35','🟢'],
        ['Poll Summary on Dashboard','Top result snippet on each event card',                          'Organiser',   'Glanceable status',               '3','3','4','2','3.05','🟢'],
        ['Delete / Leave Event',     'Organiser deletes for all; invitee removes from dashboard',      'All',         'Event lifecycle management',      '3','3','5','1','3.05','🟢'],
        ['Hard Deadline & Auto-Lock','Poll closes at set time; winning option auto-selected',          'Organiser/All','Stop procrastination',          '5','4','2','5','4.05','🟡'],
        ['Explicit Poll Finalization','Organiser taps Finalize to mark poll as decided',               'Organiser',   'Clear "settled" signal',          '5','4','4','5','4.40','🟡'],
        ['Attendee List / Who\'s Going','Show named participants who have joined/voted',               'All',         'Social proof; see who\'s in',     '4','3','3','4','3.45','🟡'],
        ['RSVP Status',              'Going / Maybe / Not Going per invitee',                         'Invitee',     'Distinct commitment signal',      '4','3','3','4','3.45','🟡'],
        ['Countdown Timer',          'Visible time remaining before poll closes',                     'Invitee',     'Urgency — drives faster votes',   '4','2','3','5','3.35','🟡'],
        ['Event Status Lifecycle',   'Collecting → Closing Soon → Completed states',                  'Invitee',     'Keeps group oriented',            '3','3','3','3','3.00','🟡'],
        ['Push Notifications',       'Alert when new poll added or deadline approaching',              'Invitee',     'Drives participation',            '4','2','2','3','2.85','🔵'],
        ['Word Parsing (LLM)',        'Parse chat messages to auto-create polls',                      'All (auto)',   'Less manual organiser setup',     '2','1','1','3','1.70','🔵'],
        ['Payment Split Integration','Venmo-style cost splitting tied to event decisions',             'Organiser',   'Convenience for cost events',     '3','1','1','2','1.80','🔵'],
    ],
    col_widths=[1.5, 1.8, 0.85, 1.2, 0.3, 0.3, 0.3, 0.3, 0.5, 0.5]
)
body('UI = User Impact · MN = MVP Necessity · IM = Implementability · LV = Learning Value')
body('')
body('Sprint 3 changes to feature analysis:', bold=True)
for b in [
    'Added "Explicit Poll Finalization" (🟡 Next Sprint) — emerged directly from user testing (3/7 users asked "is this the final answer?")',
    'Promoted "Attendee List / Who\'s Going" from 🔵 Stretch to 🟡 Next Sprint — 4/7 users requested this in testing',
    'Word Parsing remains 🔵 Stretch only — privacy concerns confirmed by Sprint 3 participants',
    'All 🟢 Built features are now validated with live Firebase data and user session observations',
]:
    bullet(b)
spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 12. NEXT LEARNING PROTOTYPE PLANS
# ══════════════════════════════════════════════════════════════════════════════
heading('12. Next Learning Prototype Plans (Sprint 4)')
hr()
sprint3_callout('Sprint 3 validated the core engagement loop. Sprint 4 focuses on the two critical gaps surfaced by testing.')

heading('Purpose', level=2)
body(
    'Sprint 3 proved that frictionless joining, real-time voting, and live results work as intended. '
    'Sprint 4 must address the two most important gaps identified:'
)
for b in [
    'The finalization problem: users do not know when a decision is truly "done." Test whether an explicit finalization mechanism creates the clear "settled" moment users expect.',
    'Participation drop-off: avg votes per poll was 2.8/~5 participants (56%). Test whether a nudge or visible participation indicator improves completion rates.',
]:
    bullet(b)

heading('Core Hypothesis for Sprint 4', level=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    'Adding an explicit poll deadline with a visible countdown will cause participants to vote sooner '
    'and feel more certain that decisions are final, without adding friction to the organiser experience.'
)
run.font.italic = True
run.font.size   = Pt(10.5)
run.font.color.rgb = MID_BLUE

heading('Features Planned for Sprint 4 Prototype', level=2)
add_table(
    ['Feature', 'What It Tests'],
    [
        ['Hard deadline per poll (organiser sets time)',    'Does deadline pressure drive faster and more complete participation?'],
        ['Explicit "Finalize Poll" button (organiser)',     'Does a manual finalization signal remove ambiguity about decision state?'],
        ['Countdown timer visible to all participants',     'Does visible urgency change voting behaviour?'],
        ['RSVP / Attendance status',                        'Is attendance tracking distinct enough from poll voting to add value without confusion?'],
        ['"Who hasn\'t voted?" indicator for organiser',   'Does showing participation gaps help organisers nudge members effectively?'],
    ],
    col_widths=[2.5, 4.8]
)

heading('Testing Plan for Sprint 4', level=2)
body('Session design: longitudinal — give real groups a real task (plan an optional event) using Polled with deadlines. Observe over 24–48 hours, not a controlled lab session.')
body('')
body('Sample: 3 groups of 4–6 participants each (12–18 total). At least one group must be an existing friend/study group, not recruited strangers.')
body('')
body('Metrics to collect:', bold=True)
for m in [
    'Time from poll creation to each vote — does deadline proximity predict vote timing?',
    'Participation rate (votes cast / participants joined) — target improvement from 56% to 75%+',
    'Comprehension: do users correctly identify the winning choice after finalization?',
    'Survey: "When did you feel the decision was made?" (open-ended)',
]:
    bullet(m)

body('')
body('Questions to answer:', bold=True)
for q in [
    'Does the hard deadline push people to respond faster? (carried over from Sprint 2 — now testable)',
    'Does the finalization state clearly communicate to all users that the decision is done?',
    'Does the attendee/RSVP feature add value or add confusion?',
    'What is the right deadline granularity — hours, days, or fully custom?',
]:
    bullet(q)

heading('What We Are NOT Building in Sprint 4', level=2)
for b in [
    'Word Parsing / LLM features — privacy concerns confirmed; insufficient ROI',
    'Push notifications — complex infrastructure; in-app indicators used as proxy',
    'Payment integrations — premature; not in core value proposition yet',
]:
    bullet(b)

spacer()
hr()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sprint 3 Notebook  ·  Team Polled (formerly Settled)  ·  Sankar Gopalkrishna · Kevin Lin · Rakshit Naidu  ·  March 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/rakshitnaidu/Documents/Polled/Sprint3_Notebook.docx'
doc.save(out)
print(f'Saved → {out}')
